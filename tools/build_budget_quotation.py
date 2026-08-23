from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "SAGIP-SJ-Year-1-Budgetary-Quotation-Draft.docx"
LOGO = ROOT / "apps" / "web" / "public" / "logo-transparent.png"

GREEN = "1F6B43"
DARK_GREEN = "164C31"
LIGHT_GREEN = "EAF4EE"
PALE_GREEN = "F5FAF7"
INK = "24312B"
MUTED = "5F6D66"
LIGHT_GRAY = "F2F4F3"
MID_GRAY = "D7DEDA"
WHITE = "FFFFFF"
AMBER = "B7791F"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, size=None, color=INK, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:cantSplit")
    tr_pr.append(element)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_cell_text(cell, *, size=9.2, color=INK, bold=False, align=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.line_spacing = 1.05
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            set_run_font(run, size=size, color=color, bold=bold)


def add_table(doc, headers, rows, widths, *, total_rows=(), amount_columns=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, label in enumerate(headers):
        hdr.cells[i].text = label
        set_cell_shading(hdr.cells[i], GREEN)
        style_cell_text(hdr.cells[i], size=9.0, color=WHITE, bold=True)
    prevent_row_split(hdr)
    for row_index, row_data in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for i, value in enumerate(row_data):
            cells[i].text = str(value)
            align = WD_ALIGN_PARAGRAPH.RIGHT if i in amount_columns else WD_ALIGN_PARAGRAPH.LEFT
            style_cell_text(cells[i], align=align)
        if row_index in total_rows:
            for cell in cells:
                set_cell_shading(cell, LIGHT_GREEN)
                style_cell_text(
                    cell,
                    size=9.3,
                    color=DARK_GREEN,
                    bold=True,
                    align=WD_ALIGN_PARAGRAPH.RIGHT if cells.index(cell) in amount_columns else WD_ALIGN_PARAGRAPH.LEFT,
                )
    set_table_geometry(table, widths)
    set_table_borders(table)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, *, color=GREEN, fill=LIGHT_GREEN):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Pt(9)
    paragraph.paragraph_format.right_indent = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.15
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "7")
    left.set(qn("w:color"), color)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    run = paragraph.add_run(f"{label} ")
    set_run_font(run, size=10.5, color=DARK_GREEN, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return paragraph


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text):
    paragraph = doc.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=INK)
    return paragraph


def add_label_paragraph(doc, label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(f"{label}: ")
    set_run_font(run, size=10.5, color=DARK_GREEN, bold=True)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return paragraph


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def set_image_alt(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_signature_table(doc):
    table = doc.add_table(rows=4, cols=2)
    values = [
        ("For the SAGIP-SJ Development Team", "For the Client"),
        ("\n\n________________________________", "\n\n________________________________"),
        ("Authorized Developer Representative", "Authorized SK/Barangay Representative"),
        ("Date: __________________________", "Date: __________________________"),
    ]
    for r, row in enumerate(values):
        for c, value in enumerate(row):
            table.cell(r, c).text = value
            style_cell_text(table.cell(r, c), size=9.5, color=INK, bold=(r == 0))
            set_cell_margins(table.cell(r, c), top=100, bottom=100, start=120, end=220)
    set_table_geometry(table, [4680, 4680])
    # Intentionally borderless form-layout override.
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    props = doc.core_properties
    props.title = "SAGIP-SJ Year 1 Budgetary Quotation"
    props.subject = "Developer-side quotation and negotiation guide for Year 1 implementation"
    props.author = "SAGIP-SJ Development Team"
    props.keywords = "SAGIP-SJ, budgetary quotation, software development, maintenance, SMS"

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(GREEN)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    h2.paragraph_format.space_before = Pt(11)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "SAGIP-SJ  |  YEAR 1 BUDGETARY QUOTATION"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    for run in header.runs:
        set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer)

    bullet_num_id = add_bullet_numbering(doc)

    # First-page proposal centerpiece.
    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_paragraph.paragraph_format.space_after = Pt(4)
    if LOGO.exists():
        shape = logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(0.75))
        set_image_alt(shape, "SAGIP-SJ logo", "SAGIP-SJ San Jose platform logo")

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(2)
    run = kicker.add_run("SAGIP-SJ")
    set_run_font(run, size=11, color=GREEN, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("BUDGETARY QUOTATION")
    set_run_font(run, size=24, color=DARK_GREEN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(15)
    run = subtitle.add_run("Year 1 Production Implementation, Security, Training, Infrastructure, SMS, and Support")
    set_run_font(run, size=11.5, color=MUTED)

    metadata = [
        ("Quotation reference", "SAGIP-SJ-BQ-2026-001"),
        ("Date", "August 23, 2026"),
        ("Prepared for", "Sangguniang Kabataan of Barangay San Jose, Rodriguez, Rizal"),
        ("Prepared by", "SAGIP-SJ Development Team"),
        ("Validity", "30 calendar days from issuance, unless extended in writing"),
        ("Status", "Budgetary quotation and negotiation guide; not yet a final contract"),
    ]
    add_table(doc, ["Quotation detail", "Value"], metadata, [2700, 6660])

    add_callout(
        doc,
        "Commercial position.",
        "The required Year 1 package is PHP 700,000. The part-time technical project manager and Android-first installable mobile application are separately selectable options. The maximum is PHP 880,000 only when both options are accepted in writing.",
    )

    doc.add_heading("1. Year 1 budget", level=1)
    summary_rows = [
        ("Acquisition and Turnover of the Existing SAGIP-SJ Competition Prototype, Including Source Code, Barangay Right-to-Use, and Compensation of the Original Project Team", "PHP 200,000"),
        ("Agreed Initial System Revisions, Production Deployment, System Hardening, and Technical Documentation", "PHP 100,000"),
        ("Independent Cybersecurity and Data Privacy Consultant for Security Assessment", "PHP 90,000"),
        ("Staff Seminars and Hands-on Training", "PHP 60,000"),
        ("Cloud Infrastructure for 12 Months", "PHP 60,000"),
        ("Annual System Maintenance, Monitoring, and Technical Support Services", "PHP 160,000"),
        ("SMS Alert Integration and Year 1 Messaging Credits", "PHP 30,000"),
        ("Base Year 1 Budget", "PHP 700,000"),
        ("Optional Part-time Technical Project Manager During Implementation", "Up to PHP 100,000"),
        ("Optional Android-first Installable Mobile Application Development", "Up to PHP 80,000"),
        ("Total Year 1 Budget Range", "PHP 700,000-PHP 880,000"),
    ]
    add_table(
        doc,
        ["Component", "Amount"],
        summary_rows,
        [7200, 2160],
        total_rows=(7, 10),
        amount_columns=(1,),
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Currency and tax treatment. ")
    set_run_font(run, size=10.5, color=DARK_GREEN, bold=True)
    run = p.add_run(
        "All amounts are in Philippine pesos and are working gross ceilings. Before issuance by the legally contracting developer or supplier, the quotation must state VAT or non-VAT status, taxes, withholding, reimbursable costs, and required supporting documents. Any tax withheld must be supported by the corresponding withholding certificate."
    )
    set_run_font(run, size=10.5, color=INK)

    doc.add_page_break()
    doc.add_heading("2. Detailed schedule of prices", level=1)
    p = doc.add_paragraph(
        "The required package is priced by deliverable or recurring service. One-time development work is not divided artificially across twelve months; payment follows the milestone and acceptance schedule in Section 4."
    )
    p.paragraph_format.space_after = Pt(8)

    price_rows = [
        ("1", "Existing prototype", "Original team compensation pool", "Upon prototype acceptance", "PHP 180,000"),
        ("1", "Existing prototype", "Repository, assets, right-to-use, and turnover records", "Upon complete turnover", "PHP 20,000"),
        ("2", "Production preparation", "Agreed initial revisions", "Accepted milestone", "PHP 45,000"),
        ("2", "Production preparation", "Production deployment and environment configuration", "Accepted milestone", "PHP 15,000"),
        ("2", "Production preparation", "Application, server, database, access, and secrets hardening", "Accepted milestone", "PHP 20,000"),
        ("2", "Production preparation", "Technical and operations documentation", "Accepted milestone", "PHP 20,000"),
        ("3", "Security consultant", "Independent assessment and engagement scoping", "Consultant milestone", "PHP 65,000"),
        ("3", "Security consultant", "Findings report and remediation workshop", "Consultant milestone", "PHP 15,000"),
        ("3", "Security consultant", "One retest and closure report", "Consultant milestone", "PHP 10,000"),
        ("4", "Staff training", "System administration seminar and exercises", "Completed session", "PHP 20,000"),
        ("4", "Staff training", "Data privacy and cybersecurity seminar and exercises", "Completed session", "PHP 20,000"),
        ("4", "Staff training", "Disaster operations simulation", "Completed session", "PHP 20,000"),
        ("5", "Cloud infrastructure", "Compute, storage, backups, monitoring, domain, DNS, and SSL/TLS", "PHP 5,000 monthly x 12", "PHP 60,000"),
        ("6", "Maintenance and support", "Maintenance, monitoring, backup checks, updates, and support", "12-month annual ceiling", "PHP 160,000"),
        ("7", "SMS", "Gateway integration, configuration, templates, targeting, logs, and testing", "Accepted milestone", "PHP 10,000"),
        ("7", "SMS", "Year 1 prepaid credits and applicable provider or sender-name charges", "Upon activation", "PHP 20,000"),
        ("", "REQUIRED PACKAGE TOTAL", "", "", "PHP 700,000"),
        ("A", "Technical Project Manager", "Part-time technical coordination for up to four months", "Up to PHP 25,000 monthly", "Up to PHP 100,000"),
        ("B", "Mobile development", "Android-first installable PWA for up to four months", "Up to PHP 20,000 monthly", "Up to PHP 80,000"),
        ("", "MAXIMUM WITH BOTH OPTIONS", "", "", "PHP 880,000"),
    ]
    add_table(
        doc,
        ["Ref.", "Category", "Specific cost item", "Billing basis", "Amount"],
        price_rows,
        [540, 1740, 3660, 1920, 1500],
        total_rows=(16, 19),
        amount_columns=(4,),
    )

    doc.add_heading("3. Required scope and acceptance evidence", level=1)
    doc.add_heading("3.1 Existing SAGIP-SJ prototype - PHP 200,000", level=2)
    add_bullet(doc, bullet_num_id, "Compensation of the original interdisciplinary project team through a separately acknowledged internal allocation.")
    add_bullet(doc, bullet_num_id, "Turnover of the existing source repository, repository history, application assets, and an inventory of prototype materials.")
    add_bullet(doc, bullet_num_id, "A documented, non-exclusive barangay right to host, operate, back up, and maintain SAGIP-SJ for official Barangay San Jose purposes after full payment.")
    add_label_paragraph(doc, "Acceptance evidence", "Repository access, turnover inventory, source and asset receipt, signed right-to-use terms, and acknowledgement of the original team compensation arrangement.")

    doc.add_heading("3.2 Initial revisions, deployment, hardening, and documentation - PHP 100,000", level=2)
    add_bullet(doc, bullet_num_id, "Implementation of only the initial revisions listed and accepted before contract signing.")
    add_bullet(doc, bullet_num_id, "Production deployment, environment configuration, HTTPS, secrets, database, upload, logging, and access-control hardening.")
    add_bullet(doc, bullet_num_id, "Architecture, database, API, deployment, backup, restoration, recovery, troubleshooting, administrator, dependency, and known-limitations documentation.")
    add_label_paragraph(doc, "Acceptance evidence", "Approved revision checklist, functioning production environment, access and recovery test records, and delivery of the technical documentation package.")

    doc.add_heading("3.3 Independent cybersecurity and privacy consultant - PHP 90,000", level=2)
    add_bullet(doc, bullet_num_id, "Independent web application, API, authorization, privilege, file-upload, server, database, backup, logging, and privacy review using staging and synthetic data.")
    add_bullet(doc, bullet_num_id, "Severity-ranked findings, a remediation workshop, one retest, and a final validation or closure report.")
    add_label_paragraph(doc, "Acceptance evidence", "Consultant engagement record, assessment report, remediation workshop record, retest results, and closure report. The evaluator must remain separate from the development team.")

    doc.add_heading("3.4 Staff seminars and hands-on training - PHP 60,000", level=2)
    add_bullet(doc, bullet_num_id, "System administration seminar and exercises.")
    add_bullet(doc, bullet_num_id, "Data privacy and cybersecurity seminar and exercises.")
    add_bullet(doc, bullet_num_id, "Disaster operations simulation and facilitated exercise.")
    add_label_paragraph(doc, "Acceptance evidence", "Training materials, quick-reference guides, attendance records, exercises, assessments, and completion report. Venue, meals, equipment, printing, and certificates are excluded unless expressly added.")

    doc.add_heading("3.5 Cloud infrastructure for 12 months - PHP 60,000", level=2)
    add_bullet(doc, bullet_num_id, "Production compute, database and file storage, encrypted off-server backups, basic monitoring, domain registration or renewal, DNS, and reasonable provider price movement within the ceiling.")
    add_bullet(doc, bullet_num_id, "SSL/TLS certificate issuance, HTTPS configuration, automatic renewal, and expiry monitoring. A no-fee automated certificate should be used unless a paid certificate is technically justified and approved.")
    add_label_paragraph(doc, "Acceptance evidence", "Provider account and billing records, running production service, working domain and HTTPS, backup verification, and monitoring access. Accounts should be held in the barangay's name where practical.")

    doc.add_heading("3.6 Maintenance, monitoring, and technical support - PHP 160,000", level=2)
    add_bullet(doc, bullet_num_id, "Correction of reproducible defects, routine dependency and security updates, monitoring, backup checks, and minor operational configuration assistance.")
    add_bullet(doc, bullet_num_id, "Up to eight support hours per month, reasonable emergency technical assistance, and quarterly health and maintenance reports.")
    add_label_paragraph(doc, "Acceptance evidence", "Monthly service record, support log, backup and monitoring checks, and quarterly report. Unused support hours do not become feature-development credit unless the final agreement states otherwise.")

    doc.add_heading("3.7 SMS alert integration and Year 1 credits - PHP 30,000", level=2)
    add_bullet(doc, bullet_num_id, "Staff-issued emergency and public-safety alert templates, recipient targeting, provider integration, delivery logs, and controlled testing.")
    add_bullet(doc, bullet_num_id, "PHP 20,000 ceiling for Year 1 prepaid message credits and applicable sender-name or provider charges; longer messages may consume multiple credits.")
    add_label_paragraph(doc, "Acceptance evidence", "Approved templates, test and production-send records, delivery-status logs, provider account ownership, credit balance or invoice records, and staff authorization controls.")

    doc.add_page_break()
    doc.add_heading("4. Billing and delivery schedule", level=1)
    add_callout(
        doc,
        "Billing rule.",
        "One-time work is billed only when the corresponding deliverable is submitted and accepted in writing. Cloud infrastructure and maintenance begin when the production service is activated, not automatically on contract signing.",
    )

    implementation_rows = [
        ("Month 1", "Prototype compensation and turnover; agreed initial revisions", "PHP 245,000", "PHP 25,000", "PHP 20,000", "PHP 290,000"),
        ("Month 2", "Deployment, hardening, documentation, and consultant assessment", "PHP 120,000", "PHP 25,000", "PHP 20,000", "PHP 165,000"),
        ("Month 3", "Consultant findings, two training sessions, and SMS integration", "PHP 65,000", "PHP 25,000", "PHP 20,000", "PHP 110,000"),
        ("Month 4", "Retest, disaster exercise, SMS credits, and implementation completion", "PHP 50,000", "PHP 25,000", "PHP 20,000", "PHP 95,000"),
        ("TOTAL", "", "PHP 480,000", "Up to PHP 100,000", "Up to PHP 80,000", "PHP 660,000"),
    ]
    add_table(
        doc,
        ["Period", "Required deliverables", "Base", "Technical PM", "Mobile", "Maximum"],
        implementation_rows,
        [900, 3480, 1260, 1260, 1140, 1320],
        total_rows=(4,),
        amount_columns=(2, 3, 4, 5),
    )

    doc.add_heading("Recurring services after production activation", level=2)
    recurring_rows = [
        ("Service Months 1-11", "PHP 5,000.00", "PHP 13,333.33", "11", "PHP 201,666.63"),
        ("Service Month 12", "PHP 5,000.00", "PHP 13,333.37", "1", "PHP 18,333.37"),
        ("YEAR TOTAL", "PHP 60,000.00", "PHP 160,000.00", "12", "PHP 220,000.00"),
    ]
    add_table(
        doc,
        ["Service period", "Cloud / month", "Support / month", "Months", "Subtotal"],
        recurring_rows,
        [2100, 1680, 1800, 1080, 2700],
        total_rows=(2,),
        amount_columns=(1, 2, 3, 4),
    )

    add_label_paragraph(doc, "Proposed payment due date", "Within 15 calendar days after written acceptance and receipt of the complete invoice or billing statement and agreed supporting documents, subject to the final signed contract.")
    add_label_paragraph(doc, "Client review period", "The client should provide one consolidated acceptance response within 10 business days of submission. Delayed feedback or access extends the schedule by at least the corresponding delay; silence is not treated as automatic acceptance unless the final contract expressly provides otherwise.")

    doc.add_heading("5. Optional components", level=1)
    doc.add_heading("5.1 Part-time Technical Project Manager - up to PHP 100,000", level=2)
    add_bullet(doc, bullet_num_id, "Up to PHP 25,000 per implementation month for a maximum of four months.")
    add_bullet(doc, bullet_num_id, "Coordinates technical scope, dependencies, deployment readiness, cloud and SMS providers, security remediation, retesting, technical risks, deliverables, documentation, and acceptance.")
    add_bullet(doc, bullet_num_id, "Does not replace the developer, the independent security evaluator, the client's acceptance authority, or a separate full-time QA function.")
    add_label_paragraph(doc, "Option activation", "This fee is payable only if a qualified external technical project manager is engaged and the option is accepted in writing.")

    doc.add_heading("5.2 Android-first Installable Mobile Application - up to PHP 80,000", level=2)
    add_bullet(doc, bullet_num_id, "Up to PHP 20,000 per implementation month for a maximum of four months.")
    add_bullet(doc, bullet_num_id, "Installable Progressive Web App using the existing SAGIP-SJ codebase, backend, authentication, and approved resident workflows.")
    add_bullet(doc, bullet_num_id, "Includes mobile flows, manifest, icons, installation guidance, device testing, release notes, and operations notes.")
    add_label_paragraph(doc, "Excluded from this option", "Separate native Android or iOS codebases, app-store publication and account fees, native push infrastructure, background location tracking, full offline synchronization, and mobile-only modules not present in the approved baseline.")

    doc.add_page_break()
    doc.add_heading("6. Commercial protections and responsibilities", level=1)

    doc.add_heading("6.1 Scope control and revisions", level=2)
    add_bullet(doc, bullet_num_id, "The quotation covers only the deliverables and limits stated in this document and the final attached revision list.")
    add_bullet(doc, bullet_num_id, "The client and developer must approve the initial revision list, acceptance criteria, dependencies, and responsible reviewers before work begins.")
    add_bullet(doc, bullet_num_id, "A new module, major workflow, redesign, integration, migration, deployment site, or material revision requires a written change request with its own price, schedule, and acceptance criteria.")
    add_bullet(doc, bullet_num_id, "Unused allowances, support hours, option ceilings, or third-party balances do not automatically convert into additional development labor or cash compensation.")

    doc.add_heading("6.2 Client dependencies and delay treatment", level=2)
    add_bullet(doc, bullet_num_id, "The client must appoint one authorized decision and acceptance contact and provide timely approvals, official content, lawful data, user lists, domain or provider access, and available staff for testing and training.")
    add_bullet(doc, bullet_num_id, "The schedule moves when client approvals, credentials, data, venues, personnel, or third-party services are delayed. The developer is not required to absorb additional work caused by those delays without a written adjustment.")
    add_bullet(doc, bullet_num_id, "After written notice, non-critical work may be rescheduled when an undisputed accepted invoice remains unpaid, subject to the final contract and applicable rules.")

    doc.add_heading("6.3 Creator rights and barangay right-to-use", level=2)
    add_bullet(doc, bullet_num_id, "No patent, registered copyright, exclusive licence, or sale of all intellectual-property rights is represented by this quotation.")
    add_bullet(doc, bullet_num_id, "After full payment, the barangay receives a non-exclusive, perpetual right to host, operate, back up, internally modify, and maintain the delivered SAGIP-SJ system for official Barangay San Jose purposes.")
    add_bullet(doc, bullet_num_id, "The creators retain authorship, attribution, reusable methods, general know-how, and pre-existing materials. Open-source and third-party components remain governed by their respective licences.")
    add_bullet(doc, bullet_num_id, "Resale, sublicensing, commercialization, or deployment for another barangay is outside the granted right unless separately agreed in writing.")

    doc.add_heading("6.4 Security, privacy, and service limitations", level=2)
    add_bullet(doc, bullet_num_id, "The developer will apply the agreed safeguards and correct accepted findings, but no system, assessment, retest, hosting service, or SMS provider can guarantee zero vulnerabilities, uninterrupted service, or delivery of every message.")
    add_bullet(doc, bullet_num_id, "Security testing must use staging and synthetic data. Production personal data may be accessed only when authorized, necessary, documented, and protected under the final data-handling terms.")
    add_bullet(doc, bullet_num_id, "Alerts remain human-issued by authorized personnel. Automated monitoring may prompt review but must not publish or send a public alert automatically.")
    add_bullet(doc, bullet_num_id, "SMS supplements official emergency communications; it is not guaranteed delivery and does not replace established emergency procedures.")

    doc.add_heading("6.5 Third-party services", level=2)
    add_bullet(doc, bullet_num_id, "Infrastructure, domain, SMS, and other provider costs are limited to the stated ceilings and should be supported by provider records where applicable.")
    add_bullet(doc, bullet_num_id, "Provider price changes, currency movements, credit expiry, sender-name approval, outages, and service-policy changes outside the developer's control may require a written adjustment or replacement provider.")
    add_bullet(doc, bullet_num_id, "The client should own the production cloud, domain, and SMS accounts where practical. Credentials must not be shared through insecure channels.")

    doc.add_heading("7. Exclusions", level=1)
    exclusions = [
        "New modules, major interface redesigns, or revisions outside the signed initial revision list.",
        "Physical siren or IoT hardware integration, device procurement, installation, or telecommunications equipment.",
        "Deployment to another barangay, public resale, commercialization, or unrestricted sublicensing.",
        "Large-scale data cleanup or migration, manual encoding, content production, or continuous data administration.",
        "Native Android or iOS development, app-store publication, full offline synchronization, or background location tracking unless separately quoted.",
        "SMS usage beyond the PHP 20,000 Year 1 credit allowance, premium messaging routes, guaranteed delivery, or unlimited recipients.",
        "Round-the-clock staffed help desk, guaranteed resolution time, unlimited support, or support beyond the stated eight-hour monthly limit.",
        "Venue, meals, equipment, printing, certificates, travel, and accommodation unless expressly included before acceptance.",
        "Paid SSL certificates unless technically justified and approved; a reputable automated no-fee certificate is the default.",
    ]
    for item in exclusions:
        add_bullet(doc, bullet_num_id, item)

    doc.add_heading("8. Quotation boundary", level=1)
    add_callout(
        doc,
        "Developer-side document.",
        "This quotation states the development team's proposed scope, price, deliverables, assumptions, and commercial protections. The recipient remains responsible for its funding source, approvals, procurement method, eligibility requirements, accounting, and audit compliance. The developer may provide reasonable technical and billing documents but does not select or warrant the recipient's procurement approach.",
        color=AMBER,
        fill="FFF8E8",
    )

    doc.add_heading("9. Selection and acknowledgement", level=1)
    selection_rows = [
        ("Required Year 1 package", "PHP 700,000", "Required"),
        ("Part-time Technical Project Manager", "Up to PHP 100,000", "Accept / Decline"),
        ("Android-first Installable Mobile Application", "Up to PHP 80,000", "Accept / Decline"),
        ("TOTAL ACCEPTED QUOTATION", "PHP __________________", "To be completed"),
    ]
    add_table(doc, ["Selection", "Amount", "Decision"], selection_rows, [5040, 2160, 2160], total_rows=(3,), amount_columns=(1,))

    p = doc.add_paragraph(
        "Acknowledgement of this budgetary quotation confirms the selected pricing scenario for negotiation. No work should begin until the parties complete the supplier details, final scope attachments, tax treatment, authorized signatures, and binding agreement or purchase document."
    )
    p.paragraph_format.space_after = Pt(14)

    add_signature_table(doc)

    doc.add_heading("Completion details before issuance", level=2)
    completion_rows = [
        ("Contracting developer/supplier legal name", "____________________________________________"),
        ("Business address", "____________________________________________"),
        ("TIN / registration / tax status", "____________________________________________"),
        ("Authorized representative and contact", "____________________________________________"),
        ("Client authorized representative", "____________________________________________"),
        ("Final quotation validity or expiry", "____________________________________________"),
        ("Approved revision-list attachment", "Attachment reference: _________________________"),
        ("Payment account and billing documents", "To be supplied through an authorized secure channel"),
    ]
    add_table(doc, ["Required issuance field", "Completion entry"], completion_rows, [3960, 5400])

    doc.add_paragraph()
    end = doc.add_paragraph("END OF BUDGETARY QUOTATION")
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(8)
    for run in end.runs:
        set_run_font(run, size=9, color=MUTED, bold=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


def validate_output(path: Path):
    # Confirm that the package is readable and every XML part is well formed.
    with ZipFile(path) as archive:
        names = archive.namelist()
        required_parts = {"[Content_Types].xml", "word/document.xml", "word/styles.xml", "word/numbering.xml"}
        missing = required_parts.difference(names)
        if missing:
            raise AssertionError(f"Missing DOCX parts: {sorted(missing)}")
        import xml.etree.ElementTree as ET

        for name in names:
            if name.endswith(".xml") or name.endswith(".rels"):
                ET.fromstring(archive.read(name))

    reopened = Document(path)
    full_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    full_text += "\n" + "\n".join(
        cell.text for table in reopened.tables for row in table.rows for cell in row.cells
    )
    required_phrases = (
        "PHP 700,000",
        "PHP 880,000",
        "Part-time Technical Project Manager",
        "SMS alert integration",
        "Cloud infrastructure for 12 months",
        "Creator rights and barangay right-to-use",
        "Developer-side document",
    )
    for phrase in required_phrases:
        if phrase not in full_text:
            raise AssertionError(f"Required quotation phrase missing: {phrase}")

    for table_number, table in enumerate(reopened.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
        if tbl_w is None or int(tbl_w.get(qn("w:w"))) != CONTENT_WIDTH_DXA:
            raise AssertionError(f"Table {table_number} has invalid width")
        if tbl_ind is None or int(tbl_ind.get(qn("w:w"))) != TABLE_INDENT_DXA:
            raise AssertionError(f"Table {table_number} has invalid indent")
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        if sum(grid_widths) != CONTENT_WIDTH_DXA:
            raise AssertionError(f"Table {table_number} grid does not total {CONTENT_WIDTH_DXA}")
        for row in table.rows:
            if len(row.cells) != len(grid_widths):
                raise AssertionError(f"Table {table_number} has inconsistent cell count")
            for index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[index]:
                    raise AssertionError(f"Table {table_number} has inconsistent cell geometry")

    if len(reopened.tables) < 8:
        raise AssertionError("Expected quotation tables are missing")
    return {"tables": len(reopened.tables), "paragraphs": len(reopened.paragraphs), "bytes": path.stat().st_size}


if __name__ == "__main__":
    built = build_document()
    print(built)
    print(validate_output(built))
